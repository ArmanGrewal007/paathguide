"""Data loader to populate the database from DOCX file."""

from docx import Document
from sqlalchemy.orm import Session

from paathguide.db import models, schemas
from paathguide.db.repository import SGGSTextRepository


class SGGSDataLoader:
    """Loader for SGGS data from DOCX file."""

    def __init__(self, db: Session):
        self.db = db
        self.repo = SGGSTextRepository(db)

    def load_from_docx_line_by_line(self, file_path: str, skip_first: int = 2) -> int:
        """
        Load SGGS data from DOCX file.

        Args:
            file_path: Path to the DOCX file
            skip_first: Number of initial lines to skip (default 2 for headers)

        Returns:
            Number of verses loaded
        """
        print(f"Loading SGGS data from {file_path}...")

        # Read document
        doc = Document(file_path)
        lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # Skip header lines
        lines = lines[skip_first:]
        print(f"Found {len(lines)} lines to process")

        # Parse and create verses
        verses_data = []
        skipped_count = 0

        for i, line in enumerate(lines):
            try:
                parsed = models.parse_verse_line(line)
                # Only add if we have valid data
                if parsed["gurmukhi_text"]:
                    verse_data = schemas.SGGSTextCreate(**parsed)
                    verses_data.append(verse_data)
                else:
                    skipped_count += 1

            except Exception as e:
                print(f"Error parsing line {i}: {line[:50]}... - {e}")
                skipped_count += 1
                continue

        print(f"Parsed {len(verses_data)} verses, skipped {skipped_count} lines")

        # Bulk insert verses
        if verses_data:
            try:
                # Insert in batches to avoid memory issues
                batch_size = 1000
                total_inserted = 0

                for i in range(0, len(verses_data), batch_size):
                    batch = verses_data[i : i + batch_size]
                    self.repo.bulk_create_sggs_texts(batch)
                    total_inserted += len(batch)
                    print(
                        f"Inserted batch {i // batch_size + 1}: {total_inserted}/{len(verses_data)} verses"
                    )

                print(f"Successfully loaded {total_inserted} verses into database")
                return total_inserted

            except Exception as e:
                print(f"Error inserting verses: {e}")
                self.db.rollback()
                raise

        return 0

    def load_by_page(self, file_path: str, skip_first: int = 2) -> int:
        """
        Load SGGS data from DOCX file, grouping all gurmukhi_text by page_number.
        Each row in the Verse table will represent one page, with all lines for that page concatenated.

        Args:
            file_path: Path to the DOCX file
            skip_first: Number of initial lines to skip (default 2 for headers)

        Returns:
            Number of pages loaded
        """
        print(f"Loading SGGS data by page from {file_path}...")

        doc = Document(file_path)
        lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        lines = lines[skip_first:]

        # Group lines by page_number
        pages = {}
        for i, line in enumerate(lines):
            try:
                parsed = models.parse_verse_line(line)
                page = parsed.get("page_number")
                text = parsed.get("gurmukhi_text", "")
                if page and text:
                    if page not in pages:
                        pages[page] = []
                    pages[page].append(text)
            except Exception as e:
                print(f"Error parsing line {i}: {line[:50]}... - {e}")
                continue

        print(f"Found {len(pages)} pages to insert")

        # Insert each page as a single SGGSText row with concatenated text and line_number=0
        inserted = 0
        for page, texts in pages.items():
            full_text = " ".join(texts)  # space-separated
            sggs_text_data = schemas.SGGSTextCreate(
                gurmukhi_text=full_text,
                page_number=page,
                line_number=0,
                translation=None,
                transliteration=None,
                raag=None,
                author=None
            )
            self.repo.create_sggs_text(sggs_text_data)
            inserted += 1

        self.db.commit()
        print(f"Inserted {inserted} pages into database (one row per page)")
        return inserted

    def clear_database(self):
        """Clear all SGGS texts from the database."""
        self.db.query(models.SGGSText).delete()
        self.db.commit()
        print("Database cleared")

    def rebuild_fts_index(self):
        """Rebuild the FTS5 index to fix search issues."""
        from sqlalchemy import text
        
        print("Rebuilding FTS5 index...")
        try:
            # Execute rebuild command
            self.db.execute(text("INSERT INTO sggs_fts(sggs_fts) VALUES('rebuild')"))
            self.db.commit()
            print("FTS5 index rebuilt successfully")
        except Exception as e:
            print(f"Error rebuilding FTS5 index: {e}")
            # If rebuild fails, try recreating the table
            print("Attempting to recreate FTS5 table...")
            models.create_fts_table()
            print("FTS5 table recreated")

    def reload_data(self, file_path: str, skip_first: int = 2) -> int:
        """Clear database and reload data."""
        self.clear_database()
        return self.load_from_docx_line_by_line(file_path, skip_first)
